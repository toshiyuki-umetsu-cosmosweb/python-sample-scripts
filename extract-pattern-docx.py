"""
補足：
    python-docxが必要。
    # pip install python-docx

使い方：
ファイル選択ダイアログを用意したり、検索パターン入力ダイアログを表示するのが面倒なので、
サクッとPythonスクリプトファイルにドラッグアンドドロップするとIDを列挙したファイルを出力する形にした。

その他：
    PDFからの抽出も考えたが、PDFの構造上パラグラフや表から抽出しにくそうだったので、
    docxから抽出する形にした。

"""
import sys
import traceback
import re
import pdb
from docx import Document

def extract_ids(ids : list[str], text : str, pattern : re.Pattern):
    """
    text中の、patternにマッチしたテキストをidsのリストに追加する。
    
    Parameters
    ----------
    ids : list[str]
        マッチしたテキストを格納するリスト
    text : str
        検索対象のテキスト
    pattern : re.Pattern
        正規表現パターン
    """
    tokens = re.findall(pattern, text)
    for token in tokens:
        if not token in ids:
            ids.append(token)

if __name__ == '__main__':
    try:
        if len(sys.argv) <= 1:
            raise Exception('File not specified.')
            
        target_file = sys.argv[1]
        
        """
        正規表現パターン。
        汎用的にどこかから持ってくるとなると、
        それ用のユーザーインタフェースを持ってくる必要があるのでやめた。
        """
        pattern_text = 'SUT-[a-zA-Z\d-]*\d'
        pattern = re.compile(pattern_text) 
        
        document = Document(target_file)
        ids = []
        
        for i,paragraph in enumerate(document.paragraphs):
            if pattern.match(paragraph.text):
                extract_ids(ids, paragraph.text, pattern)

        
        # Note: ドキュメント形式によっては,
        # テーブルからセルテキストの完全一致で検索した方が抽出しやすいかも。
        for i,table in enumerate(document.tables):
            for j,row in enumerate(table.rows):
                for k,column in enumerate(row.cells):
                    extract_ids(ids, column.text, pattern)

        # 結果を出力する。
        ids.sort()
        
        with open('ids.txt', mode="w") as f:
            for id in ids:
                f.write(id + '\n')
                print(id) # 確認のため印字

    except Exception as e:
        print(traceback.format_exc())
        print(e)
        
